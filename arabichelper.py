import eel
import os
import requests
import subprocess
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from threading import Thread


class console_log:
    logLevel = 3 # 3: debug, 2: warn, 1: error

    @staticmethod
    def debug(msg):
        if console_log.logLevel >= 3:
            print("[DEBUG]   " + msg)
    
    @staticmethod
    def warn(msg):
        if console_log.logLevel >= 2:
            print("[WARNING] " + msg)

    @staticmethod
    def error(msg):
        if console_log.logLevel >= 1:
            print("[ERROR]   " + msg)


class Word():
    def __init__(self, word):
        self.word = word
        self.meanings = []
        self.lookup_done = False

    def get_url(self):
        return "https://www.almaany.com/ar/dict/ar-ar/" + self.word


    def get_page_html(self, url):
        page = requests.get(url)
        if page:
            if page.status_code < 300:
                return page


    def get_soup(self, page):
        soup = BeautifulSoup(page.content, "html.parser")
        return soup


    def get_meanings(self, soup):
        wordlist = soup.find("ol", {"class": "meaning-results"}).findAll("li", recursive=False)
        console_log.debug("Adding meanings for word: " + self.word)
        for word in wordlist:
            self.meanings.append(word.text)


    def select_meaning(self):
        console_log.debug("Selecting meanings for word: " + self.word)
        for word in self.meanings:
            eel.add_meaning(word.text)
        eel.print_meaning()


    def lookup(self):
        url = self.get_url()
        page = self.get_page_html(url)
        if page == None:
            console_log.error("Could not get page for word: " + self.word + ", please check your internet connection")
        soup = self.get_soup(page)
        self.get_meanings(soup)
        self.lookup_done = True


def get_words():
    global WordList
    for word in WordList:
        word.lookup()


@eel.expose
def lookup_words(text):
    console_log.debug("lookup_words called")
    global WordList
    WordList = []
    words = text.splitlines()
    for word in words:
        word = word.strip()
        if word != "":
            WordList.append(Word(word))  # Word("كشافة").lookup()
    thread = Thread(target = get_words)
    thread.start()


@eel.expose
def request_words():
    console_log.debug("Requested Words")
    global WordList
    words = []
    for word in WordList:
        words.append(word.word)
    return words


@eel.expose
def get_meaning(id):
    global WordList
    if id < len(WordList):
        while True:
            if WordList[id].lookup_done:
                eel.add_meaning(WordList[id].meanings)
                break
    else:
        console_log.error("Cannot get word id: " + id + ", as it exceeds WordList length")


@eel.expose
def generate_doc(selectionList):
    global WordList
    if len(selectionList) == len(WordList):
        console_log.debug("Generating word document with selection: " + str(selectionList))
        document = Document()
        word_style = document.styles.add_style('wordstyle', WD_STYLE_TYPE.CHARACTER)
        pof_style = document.styles.add_style('pofstyle', WD_STYLE_TYPE.CHARACTER)
        meaning_style = document.styles.add_style('meaningstyle', WD_STYLE_TYPE.CHARACTER)
        i = 0
        for selection in selectionList:
            meaning_text = WordList[i].meanings[selection]
            meaning_text = meaning_text.strip()
            meaning_text = list(filter(bool, meaning_text.splitlines()))
            word_text = meaning_text[0][0:meaning_text[0].find(":")+1]
            pos_text = meaning_text[0][meaning_text[0].find(":")+1:]
            meaning_text = meaning_text[1]

            heading = document.add_heading(level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            word_run = heading.add_run(word_text)
            word_run.style = word_style
            word_font = word_run.font
            word_font.rtl = True
            word_font.size = Pt(16)
            word_font.bold = True
            word_font.color.rgb = RGBColor(0x0, 0x70, 0xC0)

            pos_run = heading.add_run(pos_text)
            pos_run.style = pof_style
            pos_font = pos_run.font
            pos_font.rtl = True
            pos_font.size = Pt(16)
            pos_font.color.rgb = RGBColor(0x0, 0x20, 0x60)

            meaning_run = paragraph.add_run(meaning_text)
            meaning_run.style = meaning_style
            meaning_font = meaning_run.font
            meaning_font.rtl = True
            i += 1
        document.save('output.docx')
        console_log.debug("Finished generating document!")
        return True
    else:
        console_log.error("SelectionList != WordList, at generate_doc function")


@eel.expose
def open_file(mode):
    file = str(os.path.dirname(os.path.abspath(__file__)) + r"\output.docx")
    if mode == 1:
        console_log.debug("Opening file in explorer")
        subprocess.Popen(r'explorer /select,"' + file + r'"')
    else:
        console_log.debug("Opening file")
        os.startfile(file)


def on_close(page, sockets):
	console_log.debug(page +  " closed")


web_options = {
	"mode": "chrome-app",
	"host": "localhost",
	"port": 8000,
    "cmdline_args": ["--disable-extensions"],
}


def main():
    print("[ Arabic Keyword Scraper Console ]\n")
    console_log.debug("Starting Program!")
    eel.init("web_view")
    try:
        eel.start("index.html", size=(507, 595), options=web_options, callback=on_close)
    except EnvironmentError:
        console_log.error("Please install Google Chrome on your system for the app to properly function")
        input("Press any key to continue...")

if __name__ == "__main__":
    main()
